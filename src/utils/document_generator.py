from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class DocumentGenerator:
    def __init__(self):
        self.downloads_folder = 'downloads'
        os.makedirs(self.downloads_folder, exist_ok=True)
    
    def create_document(self, content_data):
        """Create a Word document from scraped content"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(content_data.get('title', 'Scraped Content'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Scraped from: {content_data.get('url', 'Unknown URL')}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Word Count: {content_data.get('word_count', 0)}")
            
            # Add line break
            doc.add_paragraph("")
            
            # Add meta description if available
            meta_desc = content_data.get('meta_description', '')
            if meta_desc:
                doc.add_heading("Description", level=1)
                doc.add_paragraph(meta_desc)
                doc.add_paragraph("")
            
            # Add headings
            headings = content_data.get('headings', [])
            if headings:
                doc.add_heading("Page Structure", level=1)
                for heading in headings:
                    level = min(heading['level'] + 1, 9)  # Word supports up to level 9
                    doc.add_heading(heading['text'], level=level)
                doc.add_paragraph("")
            
            # Add main content
            if content_data.get('paragraphs'):
                doc.add_heading("Content", level=1)
                for paragraph in content_data['paragraphs']:
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add lists
            lists = content_data.get('lists', [])
            if lists:
                doc.add_heading("Lists", level=1)
                for list_item in lists:
                    list_type = list_item['type']
                    doc.add_paragraph(f"{list_type.title()} List:")
                    for item in list_item['items']:
                        if list_type == 'ordered':
                            doc.add_paragraph(item, style='List Number')
                        else:
                            doc.add_paragraph(item, style='List Bullet')
                    doc.add_paragraph("")
            
            # Add tables
            tables = content_data.get('tables', [])
            if tables:
                doc.add_heading("Tables", level=1)
                for i, table_data in enumerate(tables):
                    rows = table_data['rows']
                    if rows:
                        # Create table
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Table Grid'
                        
                        # Fill table
                        for row_idx, row in enumerate(rows):
                            for col_idx, cell_text in enumerate(row):
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text
                                
                                # Bold header row
                                if row_idx == 0 and table_data.get('has_header', False):
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                        
                        doc.add_paragraph("")
            
            # Add full text section
            full_text = content_data.get('full_text', '')
            if full_text and not content_data.get('paragraphs'):
                doc.add_heading("Full Text Content", level=1)
                # Split long text into paragraphs
                sentences = full_text.split('. ')
                current_paragraph = ""
                
                for sentence in sentences:
                    if len(current_paragraph + sentence) > 500:  # Max paragraph length
                        if current_paragraph:
                            doc.add_paragraph(current_paragraph.strip())
                        current_paragraph = sentence + '. '
                    else:
                        current_paragraph += sentence + '. '
                
                if current_paragraph:
                    doc.add_paragraph(current_paragraph.strip())
            
            # Save document
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"scraped_content_{timestamp}.docx"
            filepath = os.path.join(self.downloads_folder, filename)
            
            doc.save(filepath)
            return filepath
        
        except Exception as e:
            raise Exception(f"Error creating document: {str(e)}")
    
    def create_simple_document(self, content_text, title="Scraped Content"):
        """Create a simple document from plain text"""
        try:
            doc = Document()
            
            # Add title
            doc.add_heading(title, 0)
            
            # Add timestamp
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")
            
            # Add content
            if content_text:
                # Split into paragraphs (every 500 characters or at double newlines)
                paragraphs = content_text.split('\n\n')
                if len(paragraphs) == 1:
                    # Split by character count if no paragraph breaks
                    paragraphs = [content_text[i:i+500] for i in range(0, len(content_text), 500)]
                
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Save document
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"content_{timestamp}.docx"
            filepath = os.path.join(self.downloads_folder, filename)
            
            doc.save(filepath)
            return filepath
        
        except Exception as e:
            raise Exception(f"Error creating simple document: {str(e)}")
